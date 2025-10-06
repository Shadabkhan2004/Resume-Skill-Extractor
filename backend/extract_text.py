import pdfplumber 
import docx
import os
import io
import PyPDF2

# def extract_text_from_pdf(file_path):
#   text = ""
#   with pdfplumber.open(file_path) as pdf:
#     for page in pdf.pages:
#       text += page.extract_text() + "\n"
#   return text

# def extract_text_from_docx(file_path):
#   doc = docx.Document(file_path)
#   text = ""
#   for para in doc.paragraphs:
#     text += para.txt + "\n"
#   return text

def extract_text_from_file(filename, content):
  if filename.endswith(".pdf"):
      reader = PyPDF2.PdfReader(io.BytesIO(content))
      return " ".join(page.extract_text() for page in reader.pages)
  elif filename.endswith(".docx"):
      doc = docx.Document(io.BytesIO(content))
      return " ".join(p.text for p in doc.paragraphs)
  else:  # txt
      return content.decode("utf-8")
  
  